"""Create a corrected manuscript DOCX with unresolved items shown in red.

The source DOCX is never overwritten. Existing figures, tables and layout are
preserved. Red text is reserved for author confirmation or submission metadata.
"""

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "manuscript" / "word" / "PAH_PDE8B_manuscript.docx"
OUTPUT = ROOT / "manuscript" / "word" / "PAH_PDE8B_manuscript_corrected_redflag_20260817.docx"
RED = RGBColor(192, 0, 0)


def copy_run_style(source, target):
    if source is not None and source._element.rPr is not None:
        target._element.insert(0, deepcopy(source._element.rPr))


def replace_paragraph(paragraph, parts):
    """Replace a paragraph with [(text, red_bool), ...], preserving base style."""
    template = paragraph.runs[0] if paragraph.runs else None
    for run in list(paragraph.runs):
        paragraph._element.remove(run._element)
    for text, is_red in parts:
        run = paragraph.add_run(text)
        copy_run_style(template, run)
        if is_red:
            run.font.color.rgb = RED
    return paragraph


def insert_after(paragraph, text, *, red=False, style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    inserted = paragraph._parent.add_paragraph()
    inserted._p.getparent().remove(inserted._p)
    new_p.getparent().replace(new_p, inserted._p)
    if style:
        inserted.style = style
    run = inserted.add_run(text)
    if red:
        run.font.color.rgb = RED
    return inserted


def insert_before(paragraph, text, *, red=False, style=None):
    inserted = paragraph._parent.add_paragraph()
    inserted._p.getparent().remove(inserted._p)
    paragraph._p.addprevious(inserted._p)
    if style:
        inserted.style = style
    run = inserted.add_run(text)
    if red:
        run.font.color.rgb = RED
    return inserted


def find_exact(doc, text):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Paragraph not found: {text}")


def find_starts(doc, text):
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith(text):
            return paragraph
    raise ValueError(f"Paragraph not found: {text}")


def set_cell_text_color(cell, red=False):
    for p in cell.paragraphs:
        for r in p.runs:
            if red:
                r.font.color.rgb = RED


def main():
    doc = Document(SOURCE)

    # Submission metadata: deliberately red because only the authors can supply it.
    title = doc.paragraphs[0]
    cursor = title
    for text in (
        "[AUTHOR NAMES AND DEGREES — TO BE CONFIRMED]",
        "[AFFILIATIONS — TO BE CONFIRMED]",
        "[CORRESPONDING AUTHOR, POSTAL ADDRESS AND EMAIL — TO BE CONFIRMED]",
        "[RUNNING TITLE — TO BE CONFIRMED]",
        "[KEYWORDS — TO BE CONFIRMED]",
        "Red text identifies information that still requires author confirmation, a direct supporting citation, or journal-specific completion.",
    ):
        cursor = insert_after(cursor, text, red=True)

    # Evidence-boundary corrections. Cohort grouping remains concise by request.
    p = find_starts(doc, "Human lung transcriptomics offers direct access")
    replace_paragraph(p, [(
        "Human lung transcriptomics offers direct access to disease-associated molecular states, but individual lung cohorts are often modest in size and differ in tissue procurement and analytical platform [13-15]. A blood transcriptomic meta-analysis also documented substantial cross-study heterogeneity and limited gene-level reproducibility [17], while broader omics reviews emphasize the same need for independent validation [18]. Direct pooling can obscure cohort effects, platform effects and nested sample structures. Single-cell RNA sequencing can resolve cellular localization, but treating cells rather than participants as replicates creates pseudoreplication. Integrative analyses should therefore separate discovery from validation and retain participants and cohorts as the inferential units.",
        False,
    )])

    p = find_starts(doc, "Smooth-muscle-cell plasticity, endothelial dysfunction")
    replace_paragraph(p, [(p.text.replace("[4-10]", "[4-10,16]"), False)])

    p = find_starts(doc, "Cyclic-nucleotide signaling provides a relevant framework")
    replace_paragraph(p, [
        ("Cyclic-nucleotide signaling provides a relevant framework for this analysis. PDE5 and downstream cGMP-PKG signaling are established components of pulmonary vascular physiology and therapy [19-23]. ", False),
        ("Prostacyclin receptor signaling is generally linked to cAMP, but a direct primary or guideline citation should be added here before submission. ", True),
        ("Crosstalk between the two branches can occur through phosphodiesterases such as PDE3, whose cAMP-hydrolysing activity is inhibited by cGMP; PDE3 and PDE5 expression is increased in experimental pulmonary hypertension [24]. PDE8B is a high-affinity cAMP-specific phosphodiesterase with established tissue-specific roles in cAMP regulation [25-28], but its relationship to human PAH lung remodeling has not been defined. We asked which genes were reproducibly altered across PAH lung cohorts, which vascular cell types expressed the resulting candidates, and whether PDE8B belonged to a stable SMC and cyclic-nucleotide program.", False),
    ])

    p = find_starts(doc, "The study followed a staged discovery")
    # Keep the main-text grouping compact; details remain in the separate audit report.
    text = p.text.replace(
        "Four bulk-lung cohorts were used for primary discovery: GSE15197 (18 PAH and 13 control subjects), GSE113439 (14 PAH and 11 controls), GSE254617 (82 PAH and 52 controls) and GSE208592 (15 PAH and 18 controls). These datasets comprised 129 PAH and 94 control subjects.",
        "Four bulk-lung cohorts (GSE15197, GSE113439, GSE254617 and GSE208592) were used for primary discovery, comprising 129 PAH and 94 control subjects.",
    )
    replace_paragraph(p, [(text, False), (" GSE208592 is a public GEO series without a linked peer-reviewed article; the authors should confirm that it remains in the primary discovery set.", True)])

    p = find_starts(doc, "The cyclic-nucleotide analysis places PDE8B")
    replace_paragraph(p, [(p.text.replace("[19-24,29-32]", "[19-24,29-33]"), False)])

    p = find_starts(doc, "The next experiments should test this specific model")
    replace_paragraph(p, [(p.text.replace("[33-36]", "[34-36]"), False)])

    # Code availability is not submission-ready until a permanent repository is supplied.
    p = find_starts(doc, "Analysis scripts and figure-generation code")
    replace_paragraph(p, [("[REPLACE BEFORE SUBMISSION: provide the public code repository URL, release/tag, archival DOI and software environment; ‘available upon reasonable request’ is not independently reproducible.]", True)])

    # Insert required end-matter sections before References.
    refs_heading = find_exact(doc, "References")
    sections = [
        ("Ethics approval and consent to participate", "[AUTHOR/INSTITUTION CONFIRMATION REQUIRED: state whether analysis of de-identified public data was exempt from review or approved by the relevant ethics body, including protocol number where applicable.]"),
        ("Consent for publication", "[AUTHOR CONFIRMATION REQUIRED: state ‘Not applicable’ or provide the required consent statement.]"),
        ("Author contributions", "[AUTHOR CONFIRMATION REQUIRED: provide CRediT roles for every author.]"),
        ("Funding", "[AUTHOR CONFIRMATION REQUIRED: list funders and grant numbers, or state that no specific funding was received.]"),
        ("Acknowledgements", "[AUTHOR CONFIRMATION REQUIRED: add acknowledgements or state ‘None’.]"),
        ("Competing interests", "[AUTHOR CONFIRMATION REQUIRED: provide the journal-compliant competing-interests declaration.]"),
        ("Supplementary information", "[AUTHOR CONFIRMATION REQUIRED: verify the final inventory and numbering of Supplementary Figures 1–5 and any supplementary tables/files before submission.]"),
    ]
    for heading, body in sections:
        insert_before(refs_heading, heading, red=True, style="Heading 1")
        insert_before(refs_heading, body, red=True)

    # Correct verified reference metadata and article numbers.
    replacements = {
        "Humbert M, Kovacs": "Humbert M, Kovacs",
        "Humbert M, Guignabert": "Humbert M, Guignabert C, Bonnet S, Dorfmüller P, Klinger JR, Nicolls MR, et al. Pathology and pathobiology of pulmonary hypertension: state of the art and research perspectives. Eur Respir J. 2019;53(1):1801887. doi: 10.1183/13993003.01887-2018. PMID: 30545970.",
        "Guignabert C, Aman": "Guignabert C, Aman J, Bonnet S, Dorfmüller P, Olschewski AJ, Pullamsetti S, et al. Pathology and pathobiology of pulmonary hypertension: current insights and future directions. Eur Respir J. 2024;64(4):2401095. doi: 10.1183/13993003.01095-2024. PMID: 39209474.",
        "Hu Y, Chi L": "Hu Y, Chi L, Kuebler WM, Goldenberg NM. Perivascular Inflammation in Pulmonary Arterial Hypertension. Cells. 2020;9(11):2338. doi: 10.3390/cells9112338. PMID: 33105588.",
        "Evans CE, Cober": "Evans CE, Cober ND, Dai Z, Stewart DJ, Zhao YY. Endothelial cells in the pathogenesis of pulmonary arterial hypertension. Eur Respir J. 2021;58(3):2003957. doi: 10.1183/13993003.03957-2020. PMID: 33509961.",
        "Abid S, Marcos": "Abid S, Marcos E, Parpaleix A, Amsellem V, Breau M, Houssaini A, et al. CCR2/CCR5-mediated macrophage-smooth muscle cell crosstalk in pulmonary hypertension. Eur Respir J. 2019;54(4):1802308. doi: 10.1183/13993003.02308-2018. PMID: 31320454.",
        "Crnkovic S, Thekkekara": "Crnkovic S, Thekkekara Puthenparampil H, Mulch S, Biasin V, Radic N, Wilhelm J, et al. Adventitial fibroblasts direct smooth muscle cell-state transition in pulmonary vascular disease. eLife. 2025;13:RP98558. doi: 10.7554/eLife.98558. PMID: 40208251.",
        "Crnkovic S, Valzano": "Crnkovic S, Valzano F, Fließer E, Gindlhuber J, Thekkekara Puthenparampil H, Basil M, et al. Single-cell transcriptomics reveals skewed cellular communication and phenotypic shift in pulmonary artery remodeling. JCI Insight. 2022;7(20):e153471. doi: 10.1172/jci.insight.153471. PMID: 36099047.",
        "Saygin D, Tabib": "Saygin D, Tabib T, Bittar HET, Valenzi E, Sembrat J, Chan SY, et al. Transcriptional profiling of lung cell populations in idiopathic pulmonary arterial hypertension. Pulm Circ. 2020;10(1):2045894020908782. doi: 10.1177/2045894020908782. PMID: 32166015.",
        "Kanthapillai P,": "Barnes H, Brown Z, Burns A, Williams T. Phosphodiesterase 5 inhibitors for pulmonary hypertension. Cochrane Database Syst Rev. 2019;1(1):CD012621. doi: 10.1002/14651858.CD012621.pub2. PMID: 30701543.",
        "Sitbon O, Boucly": "Sitbon O, Boucly A, Weatherald J, Antigny F, Guignabert C, Jevnikar M, et al. Drugs targeting novel pathways in pulmonary arterial hypertension. Eur Respir J. 2025;66(2):2401830. doi: 10.1183/13993003.01830-2024. PMID: 40341051.",
    }
    for paragraph in doc.paragraphs:
        for prefix, replacement in replacements.items():
            if paragraph.text.startswith(prefix):
                if prefix == "Humbert M, Kovacs":
                    replacement = paragraph.text + " Corrigendum: Eur Heart J. 2023;44(15):1312. doi: 10.1093/eurheartj/ehad005. PMID: 36821743."
                replace_paragraph(paragraph, [(replacement, False)])
                break

    # Add a short red note to Table 1 without expanding main-text etiologic detail.
    if doc.tables:
        table = doc.tables[0]
        note_row = table.add_row()
        merged = note_row.cells[0]
        for cell in note_row.cells[1:]:
            merged = merged.merge(cell)
        merged.text = "AUTHOR CHECK: confirm final cohort counts against the locked analysis manifest before submission; detailed etiology breakdown is retained in the audit record rather than the main text."
        set_cell_text_color(merged, red=True)

    # Standardize red placeholder font size so it remains legible but unobtrusive.
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.color.rgb == RED and run.font.size is None:
                run.font.size = Pt(9.5)

    doc.core_properties.title = "PAH PDE8B manuscript — corrected red-flag review version"
    doc.core_properties.comments = "Corrections integrated 2026-08-17; red text requires author confirmation or additional support."
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
