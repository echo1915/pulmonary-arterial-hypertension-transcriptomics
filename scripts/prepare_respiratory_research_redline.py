from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


SOURCE = Path(
    "manuscript/word/PAH_PDE8B_manuscript_20260820_author_contributions_refs_ordered.docx"
)
OUTPUT = Path(
    "manuscript/word/PAH_PDE8B_manuscript_20260820_Respiratory_Research_redline.docx"
)

RED = RGBColor(0xFF, 0x00, 0x00)


def set_red(run) -> None:
    run.font.color.rgb = RED


def clear_paragraph(paragraph) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def add_red_action(paragraph, text: str) -> None:
    run = paragraph.add_run(text)
    set_red(run)
    run.bold = True


def insert_paragraph_before(paragraph, text: str):
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_paragraph = paragraph._parent.add_paragraph()
    new_paragraph._p.getparent().remove(new_paragraph._p)
    new_p.addnext(new_paragraph._p)
    new_paragraph._p.getparent().remove(new_paragraph._p)
    paragraph._p.addprevious(new_paragraph._p)
    new_paragraph.add_run(text)
    return new_paragraph


def add_page_number(paragraph) -> None:
    paragraph.alignment = 2
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instr, separate, display, end):
        run._r.append(element)


def add_continuous_line_numbering(section) -> None:
    sect_pr = section._sectPr
    for existing in sect_pr.findall(qn("w:lnNumType")):
        sect_pr.remove(existing)
    line_numbers = OxmlElement("w:lnNumType")
    line_numbers.set(qn("w:countBy"), "1")
    line_numbers.set(qn("w:start"), "1")
    line_numbers.set(qn("w:restart"), "continuous")
    sect_pr.append(line_numbers)


def remove_manual_page_breaks(document: Document) -> int:
    removed = 0
    for paragraph in document.paragraphs:
        p_pr = paragraph._p.pPr
        if p_pr is not None:
            for node in p_pr.findall(qn("w:pageBreakBefore")):
                p_pr.remove(node)
                removed += 1
        for run in paragraph.runs:
            for br in list(run._r.findall(qn("w:br"))):
                if br.get(qn("w:type")) == "page":
                    run._r.remove(br)
                    removed += 1
    return removed


def all_paragraphs(document: Document):
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    for section in document.sections:
        yield from section.header.paragraphs
        yield from section.footer.paragraphs


def strip_table_shading(document: Document) -> int:
    removed = 0
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                tc_pr = cell._tc.get_or_add_tcPr()
                for shading in tc_pr.findall(qn("w:shd")):
                    tc_pr.remove(shading)
                    removed += 1
    return removed


def add_action_before(paragraph, text: str) -> None:
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    from docx.text.paragraph import Paragraph

    action = Paragraph(new_p, paragraph._parent)
    action.style = document.styles["Normal"]
    add_red_action(action, text)


document = Document(SOURCE)

# Respiratory Research general manuscript formatting.
for paragraph in all_paragraphs(document):
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    paragraph.paragraph_format.line_spacing = 2

for section in document.sections:
    add_continuous_line_numbering(section)
    footer = section.footer
    footer.is_linked_to_previous = False
    footer_paragraph = footer.paragraphs[0]
    clear_paragraph(footer_paragraph)
    add_page_number(footer_paragraph)

manual_breaks_removed = remove_manual_page_breaks(document)
table_shading_removed = strip_table_shading(document)

# Contribution wording follows the responsibility categories shown by the user,
# adapted only to the four named authors and their authorship positions.
contrib_index = next(
    i for i, p in enumerate(document.paragraphs)
    if p.text.strip() == "Authors' contributions"
)
contrib = document.paragraphs[contrib_index + 1]
clear_paragraph(contrib)
contrib.add_run(
    "Substantial contributions to the conception or design of the work: X.L. and "
    "J.Z.; the acquisition and analysis of data: Y.C. and Q.W.; interpretation of "
    "data for the work: Y.C., Q.W., X.L. and J.Z.; drafting the work or reviewing "
    "it critically for important intellectual content: Y.C., Q.W., X.L. and J.Z.; "
    "final approval of the version to be published: all authors; agreement to be "
    "accountable for all aspects of the work in ensuring that questions related to "
    "the accuracy or integrity of any part of the work are appropriately investigated "
    "and resolved: all authors. Y.C. and Q.W. contributed equally to this work. "
)
add_red_action(
    contrib,
    "[AUTHOR CONFIRMATION REQUIRED: All four authors must confirm that these roles "
    "accurately reflect their contributions and approve the final manuscript.]",
)

# Author-supplied declarations that cannot be inferred from the manuscript.
funding_index = next(
    i for i, p in enumerate(document.paragraphs) if p.text.strip() == "Funding"
)
funding = document.paragraphs[funding_index + 1]
clear_paragraph(funding)
add_red_action(
    funding,
    "[AUTHOR ACTION REQUIRED: List every funding body and grant number, including "
    "the recipient initials; if there was no specific funding, replace this note "
    "with: 'Not applicable.']",
)

ack_index = next(
    i for i, p in enumerate(document.paragraphs) if p.text.strip() == "Acknowledgements"
)
ack = document.paragraphs[ack_index + 1]
clear_paragraph(ack)
add_red_action(
    ack,
    "[AUTHOR ACTION REQUIRED: Name non-author contributors and any writing/editorial "
    "assistance with permission; if none, replace this note with: 'Not applicable.']",
)

competing_index = next(
    i for i, p in enumerate(document.paragraphs)
    if p.text.strip() == "Competing interests"
)
competing = document.paragraphs[competing_index + 1]
competing.text = "The authors declare that they have no competing interests."
add_red_action(
    competing,
    " [AUTHOR CONFIRMATION REQUIRED: Confirm this statement for every author.]",
)

data_index = next(
    i for i, p in enumerate(document.paragraphs)
    if p.text.strip() == "Availability of data and materials"
)
data_p = document.paragraphs[data_index + 1]
data_p.add_run(" ")
add_red_action(
    data_p,
    "[AUTHOR ACTION REQUIRED: Deposit the derived machine-readable result tables "
    "and provide their persistent repository identifier or DOI before submission.]",
)

code_index = next(
    i for i, p in enumerate(document.paragraphs) if p.text.strip() == "Code availability"
)
code_p = document.paragraphs[code_index + 1]
clear_paragraph(code_p)
code_p.add_run(
    "Analysis and figure-generation scripts, environment information, path "
    "configuration and reproduction instructions will be available from "
)
add_red_action(code_p, "[PUBLIC REPOSITORY URL]")
code_p.add_run(" and archived as a versioned release at ")
add_red_action(code_p, "[ARCHIVED RELEASE DOI]")
code_p.add_run(
    ". No restricted or individual-level clinical data will be included in the release."
)

# Match the Research article section label used by the journal.
conclusion_heading = next(
    p for p in document.paragraphs if p.style.name == "Heading 1" and p.text.strip() == "Conclusion"
)
conclusion_heading.text = "Conclusions"

# Submission-system actions are flagged next to the relevant manuscript objects.
figure_caption = next(
    p for p in document.paragraphs
    if p.style.name == "Figure Caption" and p.text.strip().startswith("Figure 1")
)
add_action_before(
    figure_caption,
    "[AUTHOR ACTION REQUIRED: Upload each main figure as a separate correctly oriented "
    "composite file in first-citation order; confirm each file is <=10 MB, approximately "
    "300 dpi at final size, closely cropped, and uses embedded fonts.]",
)

supp_index = next(
    i for i, p in enumerate(document.paragraphs)
    if p.text.strip() == "Supplementary Figures"
)
supp_note = document.paragraphs[supp_index + 1]
clear_paragraph(supp_note)
add_red_action(
    supp_note,
    "[AUTHOR ACTION REQUIRED: Upload supplementary material as sequentially named "
    "'Additional file' files and provide, for each file, its filename, format, title "
    "and description; verify every additional file is cited in sequence in the text.]",
)

# Ensure the red action paragraphs also comply with double spacing.
for paragraph in all_paragraphs(document):
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    paragraph.paragraph_format.line_spacing = 2

document.save(OUTPUT)

# Structural verification.
check = Document(OUTPUT)
red_text = []
for paragraph in check.paragraphs:
    for run in paragraph.runs:
        if run.font.color.rgb == RED:
            red_text.append(run.text)

abstract_paragraphs = check.paragraphs[11:15]
abstract_words = len(re.findall(r"\b[\w'-]+\b", " ".join(p.text for p in abstract_paragraphs)))
reference_heading_index = next(
    i for i, p in enumerate(check.paragraphs) if p.text.strip() == "References"
)
references = []
for paragraph in check.paragraphs[reference_heading_index + 1:]:
    if paragraph.style.name == "List Number":
        references.append(paragraph.text)
    elif references:
        break

assert len(red_text) >= 9, red_text
assert len(references) == 36
assert all(p.paragraph_format.line_spacing == 2 for p in check.paragraphs if p.text)

# Citation order and object-caption limits.
citation_re = re.compile(r"\[([0-9,;\-–—\s]+)\]")
first_seen: list[int] = []
seen: set[int] = set()
for paragraph in check.paragraphs[:reference_heading_index]:
    for match in citation_re.finditer(paragraph.text):
        for part in re.split(r"[,;]", match.group(1)):
            part = part.strip().replace("–", "-").replace("—", "-")
            if not part:
                continue
            if "-" in part:
                start, end = map(int, part.split("-", 1))
                values = range(start, end + 1)
            else:
                values = [int(part)]
            for value in values:
                if value not in seen:
                    seen.add(value)
                    first_seen.append(value)
assert first_seen == list(range(1, 37)), first_seen

figure_caption_stats = []
for paragraph in check.paragraphs:
    if paragraph.style.name == "Figure Caption":
        title, _, legend = paragraph.text.partition("|")
        figure_caption_stats.append(
            (title.strip(), len(re.findall(r"\b[\w'-]+\b", title)),
             len(re.findall(r"\b[\w'-]+\b", legend)))
        )
assert all(title_words <= 15 for _, title_words, _ in figure_caption_stats)
assert all(legend_words <= 300 for _, _, legend_words in figure_caption_stats)

doc_xml = check._element.xml
assert 'w:type="page"' not in doc_xml
assert all(section._sectPr.find(qn("w:lnNumType")) is not None for section in check.sections)

print(f"Saved: {OUTPUT.resolve()}")
print(f"Abstract words (including labels): {abstract_words}")
print(f"References: {len(references)}")
print("Sequential citation first appearance verified: 1-36")
print(f"Figure captions checked: {len(figure_caption_stats)} (titles <=15 words; legends <=300 words)")
print(f"Red action items/runs: {len(red_text)}")
print(f"Manual page breaks removed: {manual_breaks_removed}")
print(f"Table cell shading elements removed: {table_shading_removed}")
