from __future__ import annotations

import re
from copy import deepcopy
from pathlib import Path

from docx import Document


SOURCE = Path("manuscript/word/PAH_PDE8B_manuscript_20260819.docx")
OUTPUT = Path("manuscript/word/PAH_PDE8B_manuscript_20260820_author_contributions_refs_ordered.docx")

CITATION_RE = re.compile(r"\[([0-9,;\-–—\s]+)\]")


def expand_citation(content: str) -> list[int]:
    numbers: list[int] = []
    for part in re.split(r"[,;]", content):
        part = part.strip().replace("–", "-").replace("—", "-")
        if not part:
            continue
        if "-" in part:
            start, end = (int(value) for value in part.split("-", 1))
            numbers.extend(range(start, end + 1))
        else:
            numbers.append(int(part))
    return numbers


def compress_citation(numbers: list[int]) -> str:
    ordered = sorted(dict.fromkeys(numbers))
    groups: list[str] = []
    start = previous = ordered[0]
    for number in ordered[1:] + [None]:
        if number is not None and number == previous + 1:
            previous = number
            continue
        groups.append(str(start) if start == previous else f"{start}-{previous}")
        if number is not None:
            start = previous = number
    return "[" + ",".join(groups) + "]"


def replace_citations(text: str, mapping: dict[int, int]) -> str:
    def replacement(match: re.Match[str]) -> str:
        old_numbers = expand_citation(match.group(1))
        return compress_citation([mapping[number] for number in old_numbers])

    return CITATION_RE.sub(replacement, text)


def replace_text_preserving_runs(paragraph, mapping: dict[int, int]) -> None:
    original = paragraph.text
    revised = replace_citations(original, mapping)
    if revised == original:
        return

    # Citations in this manuscript are contained within individual runs. Preserve
    # all local formatting by replacing only the text in matching runs.
    for run in paragraph.runs:
        run.text = replace_citations(run.text, mapping)
    if paragraph.text == revised:
        return

    # Defensive fallback for a citation split across Word runs.
    if not paragraph.runs:
        paragraph.add_run(revised)
        return
    paragraph.runs[0].text = revised
    for run in paragraph.runs[1:]:
        run.text = ""


def main() -> None:
    document = Document(SOURCE)
    paragraphs = document.paragraphs

    references_heading_index = next(
        index for index, paragraph in enumerate(paragraphs)
        if paragraph.text.strip() == "References"
    )

    first_appearance: list[int] = []
    seen: set[int] = set()
    for paragraph in paragraphs[:references_heading_index]:
        for match in CITATION_RE.finditer(paragraph.text):
            for number in expand_citation(match.group(1)):
                if number not in seen:
                    seen.add(number)
                    first_appearance.append(number)

    reference_paragraphs = []
    for paragraph in paragraphs[references_heading_index + 1:]:
        if paragraph.style.name == "List Number":
            reference_paragraphs.append(paragraph)
        elif reference_paragraphs:
            break

    expected = set(range(1, len(reference_paragraphs) + 1))
    if set(first_appearance) != expected:
        raise ValueError(
            f"Citation/reference mismatch: cited={sorted(seen)}, expected={sorted(expected)}"
        )

    old_to_new = {
        old_number: new_number
        for new_number, old_number in enumerate(first_appearance, start=1)
    }
    for paragraph in paragraphs[:references_heading_index]:
        replace_text_preserving_runs(paragraph, old_to_new)

    # Reorder the existing numbered reference paragraphs without recreating them,
    # which preserves their list properties and run-level formatting.
    parent = reference_paragraphs[0]._p.getparent()
    insertion_index = parent.index(reference_paragraphs[0]._p)
    reordered_elements = [
        deepcopy(reference_paragraphs[old_number - 1]._p)
        for old_number in first_appearance
    ]
    for paragraph in reference_paragraphs:
        parent.remove(paragraph._p)
    for offset, element in enumerate(reordered_elements):
        parent.insert(insertion_index + offset, element)

    contribution_heading_index = next(
        index for index, paragraph in enumerate(document.paragraphs)
        if paragraph.text.strip() == "Authors' contributions"
    )
    contribution_paragraph = document.paragraphs[contribution_heading_index + 1]
    contribution_paragraph.text = (
        "Yihang Chen and Qi Wang contributed equally to this work. Yihang Chen: "
        "methodology, software, formal analysis, data curation, visualization, and "
        "writing - original draft. Qi Wang: methodology, validation, investigation, "
        "data curation, visualization, and writing - original draft. Xiaoyan Liu: "
        "conceptualization, supervision, project administration, validation, and "
        "writing - review and editing. Jiuchang Zhong: conceptualization, supervision, "
        "project administration, and writing - review and editing. All authors read "
        "and approved the final manuscript."
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)

    # Reopen the written artifact and verify the user-facing citation sequence.
    check = Document(OUTPUT)
    check_reference_index = next(
        index for index, paragraph in enumerate(check.paragraphs)
        if paragraph.text.strip() == "References"
    )
    written_first_appearance: list[int] = []
    written_seen: set[int] = set()
    for paragraph in check.paragraphs[:check_reference_index]:
        for match in CITATION_RE.finditer(paragraph.text):
            for number in expand_citation(match.group(1)):
                if number not in written_seen:
                    written_seen.add(number)
                    written_first_appearance.append(number)
    required_sequence = list(range(1, len(reference_paragraphs) + 1))
    if written_first_appearance != required_sequence:
        raise ValueError(
            "Written citation order is not sequential: "
            f"{written_first_appearance}"
        )

    print(f"Saved: {OUTPUT.resolve()}")
    print(f"References reordered: {len(reference_paragraphs)}")
    print(f"Sequential first appearance verified: 1-{len(reference_paragraphs)}")
    print("Old-to-new mapping:", old_to_new)


if __name__ == "__main__":
    main()
