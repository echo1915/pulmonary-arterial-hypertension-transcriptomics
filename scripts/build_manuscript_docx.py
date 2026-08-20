"""Build an illustrated Word manuscript from the v0.1 Markdown draft."""

from __future__ import annotations

import csv
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from project_paths import PROJECT_DATA_ROOT


ROOT = Path(__file__).resolve().parents[1]
SUPERVISOR_MODE = "--supervisor" in sys.argv
HUMANIZED_MODE = "--humanized" in sys.argv
RESPIRATORY_RED_MODE = "--respiratory-red" in sys.argv
SOURCE = ROOT / "manuscript" / ("PAH_PDE8B_manuscript_supervisor_review.md" if SUPERVISOR_MODE else "PAH_PDE8B_manuscript_v0.1.md")
OUTDIR = ROOT / "manuscript" / "word"
OUTPUT = OUTDIR / (
    "PAH_PDE8B_manuscript_academic_humanized.docx"
    if HUMANIZED_MODE
    else (
        "PAH_PDE8B_manuscript_Respiratory_Research_author_fill_red.docx"
        if RESPIRATORY_RED_MODE
        else ("PAH_PDE8B_manuscript_supervisor_review.docx" if SUPERVISOR_MODE else "PAH_PDE8B_manuscript_with_figures_v0.1.docx")
    )
)
FIGROOT = PROJECT_DATA_ROOT / "outputs"

FIGURES = {
    "Figure 1": FIGROOT / "figure1-publication-v4" / "Figure_1_overall_project_workflow_v5_draft.png",
    "Figure 2": FIGROOT / "manuscript-figures-v3" / "Figure_2_gene_set_and_independent_validation_v6_draft.png",
    "Figure 3": FIGROOT / "figure3-localization-review-v1" / "Figure_3_cell_atlas_markers_PDE8B_SMC_review_v3.png",
    "Figure 4": FIGROOT / "figure4-smc-main-supp-review-v1" / "Figure_4_SMC_state_continuum_main_review_v1.png",
    "Figure 5": FIGROOT / "figure5-program-review-v2" / "Figure_5_PDE8B_program_main_unified_v3.png",
    "Figure 6": FIGROOT / "figure6-cyclic-nucleotide-program" / "Figure_6_PDE8B_cyclic_nucleotide_program.png",
}

SUPPS = {
    "Supplementary Figure 1": FIGROOT / "candidate-comparison" / "Supplementary_Figure_candidate_prioritization_bridge.png",
    "Supplementary Figure 2": FIGROOT / "manuscript-figures-v3" / "Supplementary_Figure_PDE8B_LOCO_v4_draft.png",
    "Supplementary Figure 3": FIGROOT / "figure4-smc-main-supp-review-v1" / "Supplementary_Figure_SMC_state_support_review_v1.png",
    "Supplementary Figure 4": FIGROOT / "figure5-program-review-v1" / "Supplementary_Figure_PDE8B_program_support_review_v1.png",
    "Supplementary Figure 5": FIGROOT / "figure5-program-review-v1" / "Supplementary_Figure_PDE8B_program_robustness_review_v1.png",
}

CAPTIONS = {
    "Figure 1": ("Overall study workflow and evidence architecture", "Four subject-level bulk-lung cohorts were integrated by cross-cohort meta-analysis to define a 188-gene robust set, followed by independent raw-CEL validation and nomination of PDE8B, PIEZO2 and SLC16A12. Two single-cell RNA-sequencing cohorts provide localization and independent replication. The convergent evidence supports an SMC-associated PDE8B/cAMP working hypothesis but does not establish causality."),
    "Figure 2": ("Cross-cohort gene-set discovery and independent lung validation", "A, Heatmap of Hedges' g values for the 188 robust genes across four discovery cohorts. B, Comparison of four-cohort meta-analytic effects with independently reprocessed GSE53408 raw-CEL effects. C, Cohort-specific effects for PDE8B, PIEZO2 and SLC16A12. D, Subject-level RNA-expression distributions in GSE53408. Discovery and validation analyses use subjects as biological replicates."),
    "Figure 3": ("Single-cell landscape, cell identity and PDE8B localization", "A, Recomputed UMAP of GSE210248 cells by annotated cell type. B, PDE8B expression on the same embedding. C, Marker dot plot supporting cell-type annotation. D,E, Patient-level pseudobulk PDE8B expression in SMC1 and SMC2 for donors (n=3) and PAH patients (n=3). Cell-level panels are descriptive; disease contrasts use patients as biological replicates."),
    "Figure 4": ("Smooth-muscle-cell state continuum and remodeling programs", "A, UMAP of SMC1 and SMC2 cells. B, Patient-pseudobulk contractile, extracellular-matrix and remodeling scores. C,D, MYH11 and FN1 along the ECM-minus-contractile state score. E, Mean paired SMC2-minus-SMC1 differences in PAH. F, Enrichment contrasting contractile/cytoskeletal and ECM/remodeling pathways. These data support state association and do not demonstrate lineage transition."),
    "Figure 5": ("PDE8B-centered vascular and cyclic-nucleotide program", "A, Cross-cohort meta-correlation with PDE8B from subject-level within-cohort correlations. B, Selected enriched pathways. C, Meta-correlations for predefined PAH and cyclic-nucleotide genes. D, Subject-level PDE8B expression in the independent GSE293580 strict-SMC subset. E, Leave-one-cohort-out program stability. F, Patient-level cell-type pseudobulk localization in GSE210248. The evidence is associative."),
    "Figure 6": ("PDE8B marks a cross-cohort cyclic-nucleotide program", "A, Network of the PDE8B-associated cyclic-nucleotide program. PDE8B hydrolyses cAMP, PDE3B provides a potential cAMP-cGMP crosstalk node, and PDE5A-PRKG1 represents the cGMP-PKG branch. B, Random-effects meta-correlations and 95% confidence intervals for prespecified genes. C, Cohort-specific correlations. D, Enrichment of the positive high-confidence PDE8B core. These data do not establish direct PDE8B regulation of cGMP or causal pathway activity."),
    "Supplementary Figure 1": ("Cross-layer prioritization of PDE8B among three validated candidates", "Parallel bulk-lung and patient-level single-cell evidence for PDE8B, PIEZO2 and SLC16A12. PDE8B was prioritized because its direction was concordant in SMCs and both independent etiologic comparisons. The comparison does not invalidate PIEZO2 or SLC16A12 and does not establish PDE8B causality."),
    "Supplementary Figure 2": ("PDE8B leave-one-cohort-out sensitivity", "Three-cohort random-effects estimates after omitting each primary cohort. Cohorts are the independent meta-analytic units."),
    "Supplementary Figure 3": ("Supporting evidence for SMC state analyses", "Patient-level and marker-level analyses supporting contractile and ECM/remodeling patterns. Inferential summaries use patient pseudobulk; cells are not independent biological replicates."),
    "Supplementary Figure 4": ("PDE8B program support and reference-gene context", "Additional pathway, predefined-gene and reference-gene analyses supporting interpretation of the PDE8B-centered coexpression program."),
    "Supplementary Figure 5": ("Robustness of the PDE8B-centered program", "Leave-one-cohort-out and related sensitivity analyses for the positive and negative PDE8B coexpression cores. These analyses assess stability but do not establish PDE8B-driven regulation."),
}

INSERT_AFTER = {
    "This structure prevented overlapping PHBI samples": "Table 1",
    "This structure prevented overlapping PHBI samples from being counted twice": "Figure 1",
    "The design avoided double-counting overlapping PHBI samples": "Table 1",
    "The design avoided double-counting overlapping PHBI samples and retained subjects": "Figure 1",
    "Subject-level GSE53408 expression retained": "Figure 2",
    "definitive disease-effect tests.": "Figure 3",
    "PDE8B initiates the transition.": "Figure 4",
    "PDE8B initiates such a transition.": "Figure 4",
    "Independent GSE293580 strict-SMC values were directionally compatible": "Figure 5",
    "direct PDE8B-mediated cGMP hydrolysis.": "Figure 6",
}


# Conservative sentence-level edits derived from academic-humanizer and sciwrite.
# These replacements deliberately leave all numbers, citations and scientific terms unchanged.
HUMANIZED_REPLACEMENTS = {
    "A clearer account of the cellular programs accompanying pulmonary vascular remodeling is therefore needed to connect human molecular observations with mechanisms that may complement established vasoactive pathways.":
        "Defining the cellular programs that accompany pulmonary vascular remodeling may connect observations in human tissue to mechanisms beyond the established vasoactive pathways.",
    "These observations create an opportunity to evaluate whether genes reproducibly altered in bulk PAH lung can be assigned to specific vascular cell types and phenotypic states.":
        "These datasets allow genes reproducibly altered in bulk PAH lung to be examined in specific vascular cell types and phenotypic states.",
    "Cyclic nucleotides provide a biologically relevant framework for such integration.":
        "Cyclic-nucleotide signaling provides a relevant framework for this analysis.",
    "We therefore asked whether a cohort-robust PAH lung signature could nominate reproducible candidates, whether single-cell evidence could resolve their cellular context, and whether PDE8B was embedded in a stable SMC and cyclic-nucleotide program.":
        "We asked three questions: which genes were reproducibly altered across PAH lung cohorts, which vascular cell types expressed the resulting candidates, and whether PDE8B belonged to a stable SMC and cyclic-nucleotide program.",
    "Here, we integrated four discovery cohorts, an independent raw-CEL validation cohort and two single-cell datasets.":
        "We integrated four discovery cohorts, an independent raw-CEL validation cohort and two single-cell datasets.",
    "Throughout, we treated these findings as association and localization evidence rather than proof that PDE8B drives PAH.":
        "We interpreted these findings as evidence of association and localization, not as proof that PDE8B drives PAH.",
    "This design reduces sensitivity to platform scale and prevents a large cohort from dominating through row-level pooling, while retaining between-cohort variation for the random-effects model.":
        "This design reduces sensitivity to platform scale, prevents a large cohort from dominating through row-level pooling and retains between-cohort variation for the random-effects model.",
    "The filtering criteria were applied jointly to prioritize reproducibility rather than maximize discovery count.":
        "We applied the filtering criteria jointly to favor reproducibility over the number of discoveries.",
    "Concordance across these roles was given more weight than raw cell count because the participant remained the inferential unit.":
        "We evaluated concordance across these roles rather than raw cell counts because participants remained the inferential units.",
    "The analysis was organized so that no validation cohort contributed to discovery (Fig. 1).":
        "No validation cohort contributed to discovery (Fig. 1).",
    "This structure prevented overlapping PHBI samples from being counted twice and maintained subjects and cohorts as the biological units.":
        "The design avoided double-counting overlapping PHBI samples and retained subjects and cohorts as the biological units.",
    "These validation estimates arose from a separately processed dataset that did not contribute to gene selection.":
        "The validation dataset was processed separately and did not contribute to gene selection.",
    "The relevant observation was preservation of direction for all three candidates, together with subject-level separation not attributable to duplicate technical columns.":
        "All three candidates retained the discovery direction, with subject-level separation that was not attributable to duplicate technical columns.",
    "This choice reflected evidence convergence rather than a claim that the other candidates were false discoveries.":
        "This prioritization reflected convergence across evidence layers; it did not imply that the other candidates were false discoveries.",
    "These data support state association but do not establish that individual cells transition from SMC1 to SMC2 or that PDE8B initiates the transition.":
        "These data associate PDE8B with SMC state, but do not show that individual cells transition from SMC1 to SMC2 or that PDE8B initiates such a transition.",
    "This stability indicated that the program was not driven by one cohort.":
        "No single cohort accounted for the stable program.",
    "Accordingly, the data defined a selective cyclic-nucleotide-associated transcriptional program rather than uniform activation of the canonical cAMP-PKA-CREB cascade.":
        "The data therefore support a selective cyclic-nucleotide-associated transcriptional program, not uniform activation of the canonical cAMP-PKA-CREB cascade.",
    "Thus, the network was not explained solely by genes sharing a disease contrast.":
        "Shared PAH-control differences alone therefore did not explain the network.",
    "The main advance is not the nomination of a single differentially expressed gene in isolation.":
        "The analysis did more than nominate an isolated differentially expressed gene.",
    "This convergence makes PDE8B a testable mechanism candidate while preserving the distinction between prioritization and causal validation.":
        "This evidence prioritizes PDE8B for mechanistic testing, but does not provide causal validation.",
    "The SMC findings refine how PDE8B should be interpreted.":
        "The SMC data constrain the interpretation of PDE8B.",
    "The cyclic-nucleotide analysis provides a mechanistic context that is relevant to PAH pharmacology.":
        "The cyclic-nucleotide analysis places PDE8B in a pathway context relevant to PAH pharmacology.",
    "The analysis also illustrates why candidate selection should remain plural until independent layers are compared.":
        "Candidate selection remained plural until the independent evidence layers had been compared.",
    "Several limitations define the boundary of the conclusions.":
        "Several limitations restrict these conclusions.",
    "The most informative next experiments should therefore test the specific model rather than broaden the computational search, particularly in view of expanding anti-remodeling therapeutic strategies in PAH [33-36].":
        "The next experiments should test this specific model rather than extend the computational search, particularly as anti-remodeling strategies in PAH expand [33-36].",
}


def humanize_lines(lines):
    if not HUMANIZED_MODE:
        return lines
    revised = []
    in_references = False
    for line in lines:
        if line.strip() == "## References":
            in_references = True
        if not in_references:
            for old, new in HUMANIZED_REPLACEMENTS.items():
                line = line.replace(old, new)
        revised.append(line)
    return revised


def set_font(run, name="Times New Roman", size=10.5, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(5)
    for style_name, size, before, after in (("Heading 1", 15, 14, 6), ("Heading 2", 12, 11, 4), ("Heading 3", 10.5, 8, 3)):
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(36, 55, 70)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    cap = styles.add_style("Figure Caption", 1)
    cap.font.name = "Times New Roman"
    cap._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    cap._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    cap.font.size = Pt(8.5)
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    cap.paragraph_format.line_spacing = 1.05
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(7)
    cap.paragraph_format.keep_together = True


def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_char1 = OxmlElement("w:fldChar"); fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar"); fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])
    set_font(run, "Times New Roman", 8, color="687782")


def add_figure(doc, label, *, new_page=False, width=5.75):
    path = FIGURES.get(label, SUPPS.get(label))
    if not path or not path.exists():
        raise FileNotFoundError(path)
    p = doc.add_paragraph()
    if new_page:
        p.paragraph_format.page_break_before = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    r = p.add_run()
    r.add_picture(str(path), width=Inches(width))
    title, body = CAPTIONS[label]
    cap = doc.add_paragraph(style="Figure Caption")
    cap.paragraph_format.keep_together = True
    lead = cap.add_run(f"{label} | {title}. ")
    set_font(lead, size=8.5, bold=True)
    rest = cap.add_run(body)
    set_font(rest, size=8.5)


def add_table1(doc):
    path = FIGROOT / "figure1-publication-v4" / "Table_1_bulk_lung_cohort_overview.csv"
    rows = list(csv.reader(path.open(encoding="utf-8-sig")))
    cap = doc.add_paragraph(style="Figure Caption")
    lead = cap.add_run("Table 1 | Overview of bulk lung cohorts. ")
    set_font(lead, size=8.5, bold=True)
    rest = cap.add_run("Cohort roles, platforms, tissue source and audited subject counts. Counts represent biological subjects rather than technical columns or individual cells.")
    set_font(rest, size=8.5)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [0.72, 0.72, 0.90, 0.53, 0.50, 0.55, 0.52, 2.05]
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            cell = table.cell(i, j)
            cell.width = Inches(widths[j])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if i == 0:
                shade_cell(cell, "DCECF4")
            elif i % 2 == 0:
                shade_cell(cell, "F4F6F7")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j in (0, 1, 3, 4, 5, 6) else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(value)
            set_font(run, "Times New Roman", 6.2 if j == 7 else 6.5, bold=(i == 0))
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))


def add_markdown_paragraph(doc, text):
    p = doc.add_paragraph()
    author_markers = (
        "AUTHOR_INPUT_NEEDED:",
        "[Author confirmation required",
        "[Author names and CRediT-aligned",
        "[AUTHOR INPUT REQUIRED",
        "will be inserted before submission",
        "will accompany the archived code release",
    )
    needs_author_input = RESPIRATORY_RED_MODE and any(
        marker.casefold() in text.casefold() for marker in author_markers
    )
    # Minimal inline handling for Markdown bold, italics and code-like placeholders.
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        bold = part.startswith("**") and part.endswith("**")
        italic = not bold and part.startswith("*") and part.endswith("*")
        code = part.startswith("`") and part.endswith("`")
        clean = part[2:-2] if bold else (part[1:-1] if italic or code else part)
        run = p.add_run(clean)
        set_font(
            run,
            "Consolas" if code else "Times New Roman",
            9 if code else 10.5,
            bold=bold,
            italic=italic,
            color="C00000" if needs_author_input else None,
        )
    return p


def add_markdown_table(doc, rows):
    """Add a manuscript table with fixed, readable geometry."""
    if not rows:
        return
    column_count = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=column_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    # Dataset audit table: favor long eligibility fields while keeping identifiers compact.
    widths = [0.72, 1.12, 0.68, 1.58, 1.78]
    if column_count != len(widths):
        widths = [6.5 / column_count] * column_count
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            cell = table.cell(i, j)
            cell.width = Inches(widths[j])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=80, start=100, bottom=80, end=100)
            if i == 0:
                shade_cell(cell, "E8EEF5")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(value)
            set_font(run, "Times New Roman", 7.2, bold=(i == 0))
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def build():
    OUTDIR.mkdir(parents=True, exist_ok=True)
    lines = humanize_lines(SOURCE.read_text(encoding="utf-8").splitlines())
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0); sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.1); sec.right_margin = Cm(2.1)
    sec.header_distance = Cm(0.8); sec.footer_distance = Cm(0.8)
    configure_styles(doc)
    add_page_number(sec)

    title = lines[0].lstrip("# ")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24); p.paragraph_format.space_after = Pt(14)
    run = p.add_run(title)
    set_font(run, "Times New Roman", 18, bold=True, color="243746")
    if RESPIRATORY_RED_MODE:
        for prompt in (
            "[AUTHOR INPUT REQUIRED: Full author names in publication order]",
            "[AUTHOR INPUT REQUIRED: Institutional affiliations with matching superscript numbers]",
            "[AUTHOR INPUT REQUIRED: Corresponding author name, postal address and email]",
        ):
            meta = doc.add_paragraph()
            meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
            meta.paragraph_format.space_after = Pt(3)
            rr = meta.add_run(prompt)
            set_font(rr, "Times New Roman", 10.5, bold=True, color="C00000")
        note = doc.add_paragraph()
        note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        note.paragraph_format.space_after = Pt(12)
        rr = note.add_run("RED TEXT REQUIRES AUTHOR COMPLETION OR FINAL VERIFICATION BEFORE SUBMISSION")
        set_font(rr, "Times New Roman", 9.5, bold=True, color="C00000")
    if not SUPERVISOR_MODE:
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.paragraph_format.space_after = Pt(18)
        r = sub.add_run("Illustrated manuscript draft v0.1")
        set_font(r, "Times New Roman", 10, italic=True, color="687782")

    inserted = set()
    previous_heading = None
    i = 1
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1; continue
        if line.startswith("## "):
            if HUMANIZED_MODE and line == previous_heading:
                i += 1
                continue
            p = doc.add_heading(line[3:], level=1)
            for run in p.runs:
                set_font(run, size=15, bold=True)
            previous_heading = line
        elif line.startswith("### "):
            p = doc.add_heading(line[4:], level=2)
            for run in p.runs:
                set_font(run, size=12, bold=True)
            previous_heading = line
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(line[2:]); set_font(run)
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            run = p.add_run(re.sub(r"^\d+\. ", "", line)); set_font(run)
        elif line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                candidate = lines[i].strip()
                cells = [cell.strip() for cell in candidate.strip("|").split("|")]
                if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
                    table_lines.append(cells)
                i += 1
            add_markdown_table(doc, table_lines)
            continue
        else:
            para = add_markdown_paragraph(doc, line)
            for marker, item in INSERT_AFTER.items():
                if marker in line and item not in inserted:
                    if item == "Table 1":
                        if not RESPIRATORY_RED_MODE:
                            add_table1(doc)
                    else:
                        add_figure(doc, item, new_page=False, width=5.75)
                    inserted.add(item)
        i += 1

    doc.add_page_break()
    p = doc.add_heading("Supplementary Figures", level=1)
    for run in p.runs:
        set_font(run, size=15, bold=True)
    for index, label in enumerate(SUPPS):
        if label not in inserted:
            add_figure(doc, label, new_page=False, width=4.85)
    doc.core_properties.title = title
    doc.core_properties.subject = "PAH transcriptomics illustrated manuscript draft"
    doc.core_properties.author = "project-001 authors"
    doc.core_properties.keywords = "PAH; PDE8B; transcriptomics; single-cell RNA sequencing"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
