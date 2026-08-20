from pathlib import Path

from docx import Document
from docx.shared import RGBColor


SOURCE = Path(
    "manuscript/word/PAH_PDE8B_manuscript_20260820_Respiratory_Research_redline.docx"
)
OUTPUT = Path(
    "manuscript/word/PAH_PDE8B_manuscript_20260820_Respiratory_Research_redline_author_roles.docx"
)
RED = RGBColor(0xFF, 0x00, 0x00)


def clear(paragraph):
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


document = Document(SOURCE)

# Yihang Chen is the sole primary/first author; remove the equal-contribution mark.
document.paragraphs[1].text = "Yihang Chen1, Qi Wang1, Xiaoyan Liu*2, Jiuchang Zhong*1"
document.paragraphs[2].text = ""

# Retain both supplied corresponding-author contacts while identifying J.Z. as primary.
document.paragraphs[8].text = (
    "*Correspondence: jczhong@mail.ccmu.edu.cn (J.Z.; primary corresponding author); "
    "lxycyyy@mail.ccmu.edu.cn (X.L.)."
)

heading_index = next(
    i for i, paragraph in enumerate(document.paragraphs)
    if paragraph.text.strip() == "Authors' contributions"
)
paragraph = document.paragraphs[heading_index + 1]
clear(paragraph)
paragraph.add_run(
    "Substantial contributions to the conception or design of the work: Y.C. and "
    "J.Z.; acquisition and analysis of data: Y.C., with support from Q.W.; "
    "interpretation of data for the work: Y.C. and J.Z., with contributions from "
    "Q.W. and X.L.; drafting the work: Y.C.; reviewing the work critically for "
    "important intellectual content: J.Z., with contributions from X.L. and Q.W.; "
    "final approval of the version to be published: all authors; agreement to be "
    "accountable for all aspects of the work in ensuring that questions related to "
    "the accuracy or integrity of any part of the work are appropriately investigated "
    "and resolved: all authors. Y.C. was the principal author and led the methodology, "
    "software implementation, formal analysis, data curation, visualization and "
    "preparation of the original draft. J.Z. was the primary corresponding author and "
    "led the study conception, supervision, project administration, interpretation of "
    "the findings and critical revision of the manuscript. "
)
confirmation = paragraph.add_run(
    "[AUTHOR CONFIRMATION REQUIRED: All authors must confirm that this statement "
    "accurately reflects their contributions and approve the final manuscript.]"
)
confirmation.font.color.rgb = RED
confirmation.bold = True

document.save(OUTPUT)

check = Document(OUTPUT)
check_heading_index = next(
    i for i, item in enumerate(check.paragraphs)
    if item.text.strip() == "Authors' contributions"
)
text = check.paragraphs[check_heading_index + 1].text
assert "Y.C. was the principal author" in text
assert "J.Z. was the primary corresponding author" in text
assert "contributed equally" not in " ".join(p.text for p in check.paragraphs[:12])
assert "primary corresponding author" in check.paragraphs[8].text

print(f"Saved: {OUTPUT.resolve()}")
print("Verified: Y.C. principal author; J.Z. primary corresponding author")
