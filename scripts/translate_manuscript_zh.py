from __future__ import annotations
import json, re, time, urllib.parse, urllib.request
from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT=Path(__file__).resolve().parents[1]
SRC=ROOT/'manuscript'/'word'/'PAH_PDE8B_manuscript_supervisor_review.docx'
OUT=ROOT/'manuscript'/'word'/'PAH_PDE8B_manuscript_supervisor_review_中文全文.docx'
CACHE=ROOT/'manuscript'/'translation_cache_zh.json'

HEADINGS={
'Abstract':'摘要','Background':'背景','Methods':'方法','Results':'结果','Discussion':'讨论','Conclusion':'结论',
'Data availability':'数据可用性','Code availability':'代码可用性','References':'参考文献','Supplementary Figures':'补充图',
'Study design and public datasets':'研究设计与公共数据集','Bulk transcriptomic preprocessing':'整体转录组数据预处理',
'Cross-cohort differential-expression meta-analysis':'跨队列差异表达荟萃分析',
'Independent validation and candidate prioritization':'独立验证与候选基因优先级排序',
'Single-cell processing, annotation and patient-level aggregation':'单细胞数据处理、注释与患者水平聚合',
'SMC-state analysis':'SMC状态分析','PDE8B cross-cohort coexpression meta-analysis':'PDE8B跨队列共表达荟萃分析',
'Functional enrichment and cyclic-nucleotide analysis':'功能富集与环核苷酸分析','Statistical analysis':'统计分析',
'A staged multi-cohort design separated discovery, validation and cellular interpretation':'分阶段多队列设计区分发现、验证与细胞层面解释',
'Cross-cohort integration identified a robust gene set and three independently validated candidates':'跨队列整合鉴定出稳健基因集及3个经独立验证的候选基因',
'Cross-layer evidence prioritized PDE8B without excluding PIEZO2 or SLC16A12':'跨层级证据优先支持PDE8B，但不排除PIEZO2或SLC16A12',
'Single-cell analysis localized PDE8B to pulmonary vascular SMC populations':'单细胞分析将PDE8B定位于肺血管SMC群体',
'PDE8B elevation accompanied a shift from contractile to remodeling-associated SMC programs':'PDE8B升高伴随SMC程序由收缩型向重塑相关状态转变',
'PDE8B marked a reproducible vascular coexpression program':'PDE8B标记了可重复的血管共表达程序',
'The PDE8B program connected cAMP-associated genes with the cGMP-PKG branch':'PDE8B程序连接cAMP相关基因与cGMP–PKG分支',
}

GLOSSARY={
'肺动脉性高血压':'肺动脉高压','平滑肌细胞':'平滑肌细胞','伪体积':'伪批量','假批量':'伪批量',
'随机效应元分析':'随机效应荟萃分析','元分析':'荟萃分析','共同表达':'共表达','队列意识':'队列感知',
'周期性核苷酸':'环核苷酸','循环核苷酸':'环核苷酸','细胞外基质（ECM）':'细胞外基质（ECM）',
'发现队列':'发现队列','验证队列':'验证队列','病因复制':'病因学重复验证','定位证据':'定位证据',
'机制模块':'机制模块','平滑肌收缩':'平滑肌收缩','错误发现率':'错误发现率',
'主成分':'主成分','单细胞RNA测序':'单细胞RNA测序','肺动脉平滑肌细胞':'肺动脉平滑肌细胞',
}

POST_REPLACE={
'大肺队列':'整体肺组织队列','大肺组':'整体肺组织队列','批量肺':'整体肺组织','肺批量':'整体肺组织',
'候选药物':'候选基因','本地化证据':'定位证据','本地化':'定位','病因复制':'病因学重复验证',
'复制队列':'重复验证队列','独立复制':'独立重复验证','发现分析':'发现阶段分析','稳健基因':'稳健基因',
'假复制':'伪重复','元相关':'荟萃相关','元 r':'荟萃 r','肺动脉高血压突破倡议':'肺高血压突破计划',
'程序分析':'基因程序分析','血管程序':'血管相关程序','合成 SMC':'合成型SMC','成纤维细胞样':'成纤维细胞样',
'受试者':'研究对象','疾病效应':'疾病相关效应','效果大小':'效应量','效果':'效应',
}

MANUAL={
'Across 15,015 genes tested in the four bulk cohorts, 543 met the criteria for the stable PDE8B mechanism module and 201 met the high-confidence core threshold of |meta-r| éˆ®?0.40 (Fig. 5a). Of the stable genes, 417 were positively and 126 negatively associated with PDE8B. Leave-one-cohort-out analysis preserved the correlation sign for all 201 core genes, and 198 retained a minimum |meta-r| of at least 0.30 after each omission (Fig. 5e). This stability indicated that the program was not driven by one cohort.':'在4个整体肺组织队列中共检验15,015个基因，其中543个基因符合稳定PDE8B机制模块的判定标准，201个基因达到高置信度核心阈值（|meta-r| ≥ 0.40；图5a）。在稳定模块中，417个基因与PDE8B呈正相关，126个基因呈负相关。逐一剔除队列分析显示，全部201个核心基因的相关方向均保持不变；每次剔除后，198个基因的最低|meta-r|仍不低于0.30（图5e）。这些结果表明，该基因程序并非由某一个队列单独驱动。',
}

def translate(text:str)->str:
    if text in MANUAL: return MANUAL[text]
    q=urllib.parse.urlencode({'client':'gtx','sl':'en','tl':'zh-CN','dt':'t','q':text})
    req=urllib.request.Request('https://translate.googleapis.com/translate_a/single?'+q,headers={'User-Agent':'Mozilla/5.0'})
    data=json.loads(urllib.request.urlopen(req,timeout=30).read().decode('utf-8'))
    out=''.join(x[0] for x in data[0] if x and x[0])
    for a,b in GLOSSARY.items(): out=out.replace(a,b)
    for a,b in POST_REPLACE.items(): out=out.replace(a,b)
    out=out.replace('图 ', '图').replace('表 ', '表').replace('补充图 ', '补充图')
    return out

def sig(s:str):
    return re.findall(r'(?:GSE\d+|PDE8B|PIEZO2|SLC16A12|[A-Z][A-Z0-9]{2,}|[-+]?\d+(?:[.,]\d+)*(?:\s*[×x]\s*10[^ ,;)]*)?)',s)

def set_run_font(run, size=None, bold=None):
    run.font.name='Times New Roman'
    rfonts=run._element.get_or_add_rPr().get_or_add_rFonts()
    rfonts.set(qn('w:ascii'),'Times New Roman'); rfonts.set(qn('w:hAnsi'),'Times New Roman')
    rfonts.set(qn('w:eastAsia'),'SimSun'); rfonts.set(qn('w:cs'),'Times New Roman')
    if size: run.font.size=Pt(size)
    if bold is not None: run.bold=bold

def replace_paragraph(p,text):
    p.clear(); r=p.add_run(text); set_run_font(r, 12 if p.style.name=='Heading 2' else (15 if p.style.name=='Heading 1' else 10.5), p.style.name.startswith('Heading'))

def main():
    cache=json.loads(CACHE.read_text(encoding='utf-8')) if CACHE.exists() else {}
    d=Document(SRC)
    seen_refs=False
    audit=[]
    for i,p in enumerate(d.paragraphs):
        t=p.text.strip()
        if not t: continue
        if t=='References':
            if seen_refs: p.clear()
            else: replace_paragraph(p,'参考文献'); seen_refs=True
            continue
        if seen_refs and p.style.name=='List Number': continue
        if t in HEADINGS: replace_paragraph(p,HEADINGS[t]); continue
        if i==0:
            replace_paragraph(p,'跨队列转录组分析揭示PDE8B与肺动脉高压平滑肌重塑及环核苷酸程序的关联'); continue
        if t not in cache:
            cache[t]=translate(t); CACHE.write_text(json.dumps(cache,ensure_ascii=False,indent=2),encoding='utf-8'); time.sleep(.15)
        if t.startswith('Across 15,015 genes tested in the four bulk cohorts'):
            z='在4个整体肺组织队列中共检验15,015个基因，其中543个基因符合稳定PDE8B机制模块的判定标准，201个基因达到高置信度核心阈值（|meta-r| ≥ 0.40；图5a）。在稳定模块中，417个基因与PDE8B呈正相关，126个基因呈负相关。逐一剔除队列分析显示，全部201个核心基因的相关方向均保持不变；每次剔除后，198个基因的最低|meta-r|仍不低于0.30（图5e）。这些结果表明，该基因程序并非由某一个队列单独驱动。'
        else:
            z=MANUAL.get(t,cache[t])
        for a,b in POST_REPLACE.items(): z=z.replace(a,b)
        # Exact preservation gate for dataset IDs, gene symbols and all Arabic-number strings.
        en_tokens=re.findall(r'GSE\d+|[A-Z][A-Z0-9]{2,}|[-+]?\d+(?:[.,]\d+)*',t)
        zh_tokens=re.findall(r'GSE\d+|[A-Z][A-Z0-9]{2,}|[-+]?\d+(?:[.,]\d+)*',z)
        missing=[x for x in en_tokens if x not in zh_tokens]
        audit.append({'paragraph':i,'missing':missing,'en':t,'zh':z})
        replace_paragraph(p,z)
    # Translate table text while preserving numeric cells and accessions.
    for table in d.tables:
        for row in table.rows:
            for cell in row.cells:
                t=cell.text.strip()
                if not t or re.fullmatch(r'[\d.,+/–-]+|GSE\d+',t): continue
                if t not in cache:
                    cache[t]=translate(t); CACHE.write_text(json.dumps(cache,ensure_ascii=False,indent=2),encoding='utf-8'); time.sleep(.15)
                for p in cell.paragraphs: replace_paragraph(p,cache[t])
    # Explicit Chinese/Latin font pairing in styles.
    for s in d.styles:
        if not hasattr(s,'font'): continue
        s.font.name='Times New Roman'
        if s._element.rPr is not None:
            rf=s._element.rPr.get_or_add_rFonts(); rf.set(qn('w:ascii'),'Times New Roman'); rf.set(qn('w:hAnsi'),'Times New Roman'); rf.set(qn('w:eastAsia'),'SimSun'); rf.set(qn('w:cs'),'Times New Roman')
    d.core_properties.title='PDE8B与肺动脉高压平滑肌重塑及环核苷酸程序：中文全文'
    d.save(OUT)
    CACHE.write_text(json.dumps(cache,ensure_ascii=False,indent=2),encoding='utf-8')
    bad=[x for x in audit if x['missing']]
    (ROOT/'manuscript'/'translation_audit_zh.json').write_text(json.dumps({'paragraphs':len(audit),'token_mismatches':bad},ensure_ascii=False,indent=2),encoding='utf-8')
    print(OUT); print('translated_paragraphs',len(audit),'token_mismatches',len(bad))

if __name__=='__main__': main()
